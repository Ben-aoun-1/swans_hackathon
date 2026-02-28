"""Generate the retainer agreement Word template with Clio merge fields.

Run: python3 backend/scripts/generate_retainer_template.py
Output: templates/retainer_agreement.docx
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


ROOT = Path(__file__).resolve().parent.parent.parent
OUTPUT = ROOT / "templates" / "retainer_agreement.docx"


def _set_run(run, bold=False, size=None, font_name="Times New Roman", color=None):
    """Apply formatting to a run."""
    run.bold = bold
    run.font.name = font_name
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_heading_text(doc, text, size=16, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Add a centered heading paragraph."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    _set_run(run, bold=bold, size=size)
    return p


def _add_body(doc, text, bold=False, size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    _set_run(run, bold=bold, size=size)
    return p


def _add_section_heading(doc, number, title):
    """Add a numbered section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{number}. {title}")
    _set_run(run, bold=True, size=12)
    return p


def _add_field_line(doc, label, merge_field):
    """Add a label: <<merge_field>> line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run_label = p.add_run(f"{label}: ")
    _set_run(run_label, bold=True, size=11)
    run_value = p.add_run(merge_field)
    _set_run(run_value, size=11)
    return p


def _add_signature_block(doc, name_field, role_label):
    """Add a signature line with name merge field."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("_" * 40 + "     Date: " + "_" * 15)
    _set_run(run, size=11)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    run2 = p2.add_run(f"{name_field} ({role_label})")
    _set_run(run2, size=11)
    return p2


def create_retainer_template():
    """Build the retainer agreement .docx with Clio merge fields."""
    doc = Document()

    # -- Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # -- Firm header
    _add_heading_text(doc, "RICHARDS & LAW", size=18, bold=True)
    _add_heading_text(doc, "Attorneys at Law", size=12, bold=False)
    _add_body(doc, "<<Firm.Address>>", size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _add_body(doc, "Phone: <<Firm.Phone>>", size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # -- Horizontal rule (thin table border)
    p_hr = doc.add_paragraph()
    p_hr.paragraph_format.space_after = Pt(12)
    run_hr = p_hr.add_run("_" * 72)
    _set_run(run_hr, size=8, color=(150, 150, 150))

    # -- Title
    _add_heading_text(doc, "RETAINER AGREEMENT", size=16, bold=True)
    _add_body(doc, "", space_after=4)

    # -- Opening
    _add_body(
        doc,
        'This Retainer Agreement ("Agreement") is entered into on <<Today>> between:',
        size=12,
        space_after=12,
    )

    # -- Attorney info
    _add_body(doc, "ATTORNEY:", bold=True, size=12, space_after=2)
    _add_body(doc, "<<Matter.ResponsibleAttorney.Name>>", size=12, space_after=2)
    _add_body(doc, "Richards & Law", size=12, space_after=2)
    _add_body(doc, "<<Firm.Address>>", size=12, space_after=2)
    _add_body(doc, "New York, NY", size=12, space_after=12)

    # -- Client info
    _add_body(doc, "CLIENT:", bold=True, size=12, space_after=2)
    _add_body(doc, "<<Matter.CustomField.Plaintiff Name>>", size=12, space_after=2)
    _add_body(doc, "<<Matter.CustomField.Plaintiff Address>>", size=12, space_after=2)
    _add_field_line(doc, "Date of Birth", "<<Matter.CustomField.Plaintiff DOB>>")
    _add_field_line(doc, "Phone", "<<Matter.CustomField.Plaintiff Phone>>")
    _add_field_line(doc, "Email", "<<Matter.Client.Email>>")
    _add_body(doc, "", space_after=8)

    # -- Matter Reference
    _add_field_line(doc, "Matter Number", "<<Matter.DisplayNumber>>")
    _add_body(doc, "", space_after=8)

    # =====================================================================
    # SECTION 1: SCOPE OF REPRESENTATION
    # =====================================================================
    _add_section_heading(doc, 1, "SCOPE OF REPRESENTATION")

    _add_body(
        doc,
        "The Client hereby retains the Attorney to represent the Client in connection "
        "with a personal injury claim arising from a motor vehicle accident that occurred "
        "on <<Matter.CustomField.Accident Date>> at <<Matter.CustomField.Accident Location>>.",
        size=12,
        space_after=8,
    )

    _add_body(doc, "Accident Description:", bold=True, size=12, space_after=2)
    _add_body(doc, "<<Matter.CustomField.Accident Description>>", size=12, space_after=8)

    _add_field_line(doc, "Police Report Number", "<<Matter.CustomField.Police Report Number>>")
    _add_body(doc, "", space_after=4)

    # =====================================================================
    # SECTION 2: PARTIES INVOLVED
    # =====================================================================
    _add_section_heading(doc, 2, "PARTIES INVOLVED")

    _add_body(doc, "Adverse Party:", bold=True, size=12, space_after=2)
    _add_field_line(doc, "Name", "<<Matter.CustomField.Defendant Name>>")
    _add_field_line(doc, "Address", "<<Matter.CustomField.Defendant Address>>")
    _add_field_line(doc, "Insurance Company", "<<Matter.CustomField.Defendant Insurance>>")
    _add_field_line(doc, "Policy Number", "<<Matter.CustomField.Defendant Policy Number>>")
    _add_field_line(doc, "Vehicle", "<<Matter.CustomField.Defendant Vehicle>>")
    _add_body(doc, "", space_after=4)

    _add_body(doc, "Client:", bold=True, size=12, space_after=2)
    _add_field_line(doc, "Vehicle", "<<Matter.CustomField.Plaintiff Vehicle>>")
    _add_field_line(doc, "Injuries Reported", "<<Matter.CustomField.Injuries Reported>>")
    _add_body(doc, "", space_after=4)

    # =====================================================================
    # SECTION 3: ATTORNEY'S FEES
    # =====================================================================
    _add_section_heading(doc, 3, "ATTORNEY'S FEES")

    _add_body(
        doc,
        "The Attorney's fee shall be calculated on a contingency basis as follows:",
        size=12,
        space_after=6,
    )

    _add_body(
        doc,
        "(a) Thirty-three and one-third percent (33\u2153%) of the gross recovery obtained "
        "on behalf of the Client, whether by settlement, judgment, or otherwise, if the "
        "case is resolved before the filing of a lawsuit.",
        size=12,
        space_after=4,
    )

    _add_body(
        doc,
        "(b) Forty percent (40%) of the gross recovery if the case proceeds to litigation "
        "and a lawsuit is filed.",
        size=12,
        space_after=4,
    )

    _add_body(
        doc,
        "If no recovery is obtained, the Client shall owe no attorney's fees. The Client "
        "remains responsible for costs and disbursements as described in Section 5 below.",
        size=12,
        space_after=4,
    )

    # =====================================================================
    # SECTION 4: STATUTE OF LIMITATIONS
    # =====================================================================
    _add_section_heading(doc, 4, "STATUTE OF LIMITATIONS")

    _add_body(
        doc,
        "The applicable statute of limitations for this claim expires on "
        "<<Matter.CustomField.Statute of Limitations Date>>. The Attorney shall take "
        "all necessary actions to preserve the Client's claims within this period. "
        "Failure to act before this date may result in the permanent loss of the "
        "Client's right to seek compensation.",
        size=12,
        space_after=4,
    )

    # =====================================================================
    # SECTION 5: COSTS AND DISBURSEMENTS
    # =====================================================================
    _add_section_heading(doc, 5, "COSTS AND DISBURSEMENTS")

    _add_body(
        doc,
        "The Attorney shall advance all costs and disbursements necessary for the "
        "prosecution of the Client's claim, including but not limited to court filing "
        "fees, medical record retrieval fees, expert witness fees, deposition costs, "
        "and investigation expenses. Such costs shall be deducted from the Client's "
        "share of any recovery obtained. If no recovery is made, the Client shall not "
        "be responsible for reimbursing the Attorney for advanced costs.",
        size=12,
        space_after=4,
    )

    # =====================================================================
    # SECTION 6: CLIENT'S OBLIGATIONS
    # =====================================================================
    _add_section_heading(doc, 6, "CLIENT'S OBLIGATIONS")

    _add_body(doc, "The Client agrees to:", bold=False, size=12, space_after=4)

    obligations = [
        "Cooperate fully with the Attorney and provide all requested information, "
        "documents, and records in a timely manner.",
        "Attend all medical appointments, depositions, court appearances, and other "
        "proceedings as reasonably required.",
        "Not discuss the case with opposing parties, their attorneys, or their insurance "
        "representatives without first consulting the Attorney.",
        "Promptly inform the Attorney of any change in contact information, medical "
        "condition, or other circumstances relevant to the case.",
        "Not settle or compromise the claim without the Attorney's written consent.",
    ]

    for i, obligation in enumerate(obligations, 1):
        _add_body(doc, f"({chr(96 + i)}) {obligation}", size=12, space_after=3)

    # =====================================================================
    # SECTION 7: ATTORNEY'S OBLIGATIONS
    # =====================================================================
    _add_section_heading(doc, 7, "ATTORNEY'S OBLIGATIONS")

    _add_body(
        doc,
        "The Attorney agrees to diligently pursue the Client's claim, keep the Client "
        "reasonably informed of the status of the case, promptly respond to the Client's "
        "inquiries, and obtain the Client's consent before accepting any settlement offer.",
        size=12,
        space_after=4,
    )

    # =====================================================================
    # SECTION 8: TERMINATION
    # =====================================================================
    _add_section_heading(doc, 8, "TERMINATION")

    _add_body(
        doc,
        "Either party may terminate this Agreement upon written notice to the other party. "
        "If the Client terminates this Agreement, the Attorney shall be entitled to "
        "compensation on a quantum meruit basis for services rendered prior to termination, "
        "to be paid from any recovery subsequently obtained. If the Attorney withdraws from "
        "representation, the Attorney shall take reasonable steps to protect the Client's "
        "interests and provide adequate notice.",
        size=12,
        space_after=4,
    )

    # =====================================================================
    # SECTION 9: DISPUTE RESOLUTION
    # =====================================================================
    _add_section_heading(doc, 9, "DISPUTE RESOLUTION")

    _add_body(
        doc,
        "Any dispute arising under this Agreement shall be resolved through arbitration "
        "in accordance with the rules of the American Arbitration Association, with the "
        "arbitration to take place in the State of New York. The decision of the arbitrator "
        "shall be final and binding on both parties.",
        size=12,
        space_after=4,
    )

    # =====================================================================
    # SECTION 10: ENTIRE AGREEMENT
    # =====================================================================
    _add_section_heading(doc, 10, "ENTIRE AGREEMENT")

    _add_body(
        doc,
        "This Agreement constitutes the entire agreement between the parties and supersedes "
        "all prior negotiations, representations, warranties, commitments, offers, and "
        "contracts of any kind, whether written or oral. This Agreement may not be modified "
        "except by a written instrument signed by both parties.",
        size=12,
        space_after=4,
    )

    # =====================================================================
    # SECTION 11: GOVERNING LAW
    # =====================================================================
    _add_section_heading(doc, 11, "GOVERNING LAW")

    _add_body(
        doc,
        "This Agreement shall be governed by and construed in accordance with the laws "
        "of the State of New York.",
        size=12,
        space_after=4,
    )

    # =====================================================================
    # ACKNOWLEDGMENT
    # =====================================================================
    _add_body(doc, "", space_after=8)
    _add_body(
        doc,
        "BY SIGNING BELOW, THE CLIENT ACKNOWLEDGES THAT THEY HAVE READ THIS AGREEMENT "
        "IN ITS ENTIRETY, UNDERSTAND ITS TERMS, AND AGREE TO BE BOUND BY THEM.",
        bold=True,
        size=11,
        space_after=12,
    )

    # -- Signature blocks
    _add_signature_block(doc, "<<Matter.CustomField.Plaintiff Name>>", "Client")
    _add_signature_block(doc, "<<Matter.ResponsibleAttorney.Name>>", "Attorney, Richards & Law")

    # -- Save
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Retainer template saved to: {OUTPUT}")


if __name__ == "__main__":
    create_retainer_template()
